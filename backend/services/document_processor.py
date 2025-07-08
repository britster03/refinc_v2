"""
Document Processing Service

Handles extraction of text from various document formats including:
- PDF files (text extraction + vision-based processing)
- DOC/DOCX files
- Plain text files
- Vision-based PDF processing for format preservation
"""

import os
import logging
import base64
import io
from typing import Dict, Any, Optional, Union, List
from pathlib import Path
import tempfile

# Document processing libraries
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import python_docx
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

# Vision processing
try:
    from PIL import Image
    import pdf2image
    VISION_AVAILABLE = True
except ImportError:
    VISION_AVAILABLE = False

from groq import Groq
import aiofiles

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Advanced document processing service with vision capabilities
    """
    
    def __init__(self, groq_client: Optional[Groq] = None):
        self.groq_client = groq_client
        self.supported_formats = {'.pdf', '.doc', '.docx', '.txt'}
        
        # Debug logging for library availability
        logger.info(f"Library availability - PDF_AVAILABLE: {PDF_AVAILABLE}, DOCX_AVAILABLE: {DOCX_AVAILABLE}, PYTHON_DOCX_AVAILABLE: {PYTHON_DOCX_AVAILABLE}, VISION_AVAILABLE: {VISION_AVAILABLE}")
        
        # Check available libraries
        self.capabilities = {
            'pdf_text': PDF_AVAILABLE,
            'docx_text': DOCX_AVAILABLE or PYTHON_DOCX_AVAILABLE,
            'pdf_vision': VISION_AVAILABLE and groq_client is not None
        }
        
        logger.info(f"Document processor initialized with capabilities: {self.capabilities}")
    
    async def process_document(
        self, 
        file_path: str, 
        use_vision: bool = False,
        preserve_formatting: bool = True
    ) -> Dict[str, Any]:
        """
        Process document and extract text with optional vision processing
        
        Args:
            file_path: Path to the document file
            use_vision: Whether to use vision capabilities for PDFs
            preserve_formatting: Whether to preserve document formatting
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        result = {
            'text': '',
            'metadata': {
                'file_path': file_path,
                'file_extension': file_ext,
                'file_size': os.path.getsize(file_path),
                'processing_method': 'unknown',
                'formatting_preserved': False,
                'vision_used': False,
                'pages_processed': 0
            },
            'success': False,
            'error': None
        }
        
        try:
            if file_ext == '.pdf':
                if use_vision and self.capabilities['pdf_vision']:
                    result = await self._process_pdf_with_vision(file_path, result)
                else:
                    result = await self._extract_pdf_text(file_path, result)
            elif file_ext in ['.doc', '.docx']:
                result = await self._extract_docx_text(file_path, result)
            elif file_ext == '.txt':
                result = await self._extract_text_file(file_path, result)
            
            result['success'] = True
            
        except Exception as e:
            logger.error(f"Document processing failed for {file_path}: {str(e)}")
            result['error'] = str(e)
            result['success'] = False
        
        return result
    
    async def _extract_pdf_text(self, file_path: str, result: Dict[str, Any]) -> Dict[str, Any]:
        """Extract text from PDF using text-based methods"""
        text_content = ""
        pages_processed = 0
        
        # Try pdfplumber first (better for complex layouts)
        if PDF_AVAILABLE:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n\n"
                        pages_processed += 1
                
                if text_content.strip():
                    result['text'] = text_content.strip()
                    result['metadata']['processing_method'] = 'pdfplumber'
                    result['metadata']['pages_processed'] = pages_processed
                    return result
                    
            except Exception as e:
                logger.warning(f"PDFPlumber failed, trying PyPDF2: {str(e)}")
        
        # Fallback to PyPDF2
        if PDF_AVAILABLE:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    pages_processed = len(pdf_reader.pages)
                    
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n\n"
                
                result['text'] = text_content.strip()
                result['metadata']['processing_method'] = 'pypdf2'
                result['metadata']['pages_processed'] = pages_processed
                
            except Exception as e:
                logger.error(f"PyPDF2 also failed: {str(e)}")
                raise Exception(f"PDF text extraction failed: {str(e)}")
        else:
            raise Exception("PDF processing libraries not available")
        
        return result
    
    async def _process_pdf_with_vision(self, file_path: str, result: Dict[str, Any]) -> Dict[str, Any]:
        """Process PDF using vision capabilities to preserve formatting"""
        if not self.capabilities['pdf_vision']:
            raise Exception("Vision processing not available")
        
        try:
            # Convert PDF pages to images
            images = pdf2image.convert_from_path(file_path, dpi=300)
            
            all_text = ""
            pages_processed = 0
            
            for i, image in enumerate(images):
                # Convert PIL Image to base64
                buffer = io.BytesIO()
                image.save(buffer, format='PNG')
                img_base64 = base64.b64encode(buffer.getvalue()).decode()
                
                # Use Groq vision model to extract text
                page_text = await self._extract_text_with_vision(img_base64, i + 1)
                if page_text:
                    all_text += f"--- Page {i + 1} ---\n{page_text}\n\n"
                pages_processed += 1
            
            result['text'] = all_text.strip()
            result['metadata']['processing_method'] = 'vision_ocr'
            result['metadata']['vision_used'] = True
            result['metadata']['formatting_preserved'] = True
            result['metadata']['pages_processed'] = pages_processed
            
        except Exception as e:
            logger.error(f"Vision processing failed: {str(e)}")
            # Fallback to text extraction
            result = await self._extract_pdf_text(file_path, result)
            result['metadata']['vision_fallback'] = True
        
        return result
    
    async def _extract_text_with_vision(self, image_base64: str, page_num: int) -> str:
        """Extract text from image using Groq vision model"""
        try:
            # Note: This is a placeholder for vision API call
            # You'll need to adjust based on Groq's vision API format
            messages = [
                {
                    "role": "system",
                    "content": "You are an expert document reader. Extract all text from this resume page, preserving the original formatting, structure, and layout as much as possible. Include headers, bullet points, and maintain the visual hierarchy."
                },
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": f"Please extract all text from this resume page {page_num}, maintaining the original formatting:"
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{image_base64}"
                            }
                        }
                    ]
                }
            ]
            
            # This would need to be adjusted based on Groq's actual vision API
            response = self.groq_client.chat.completions.create(
                model="llava-v1.5-7b-4096-preview",  # Adjust model name as needed
                messages=messages,
                max_tokens=2000,
                temperature=0.1
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            logger.error(f"Vision text extraction failed for page {page_num}: {str(e)}")
            return ""
    
    async def _extract_docx_text(self, file_path: str, result: Dict[str, Any]) -> Dict[str, Any]:
        """Extract text from DOCX files"""
        if not self.capabilities['docx_text']:
            raise Exception("DOCX processing libraries not available")
        
        try:
            doc = Document(file_path)
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text_content += " | ".join(row_text) + "\n"
            
            result['text'] = text_content.strip()
            result['metadata']['processing_method'] = 'python_docx'
            result['metadata']['paragraphs'] = len(doc.paragraphs)
            result['metadata']['tables'] = len(doc.tables)
            
        except Exception as e:
            logger.error(f"DOCX processing failed: {str(e)}")
            raise Exception(f"DOCX text extraction failed: {str(e)}")
        
        return result
    
    async def _extract_text_file(self, file_path: str, result: Dict[str, Any]) -> Dict[str, Any]:
        """Extract text from plain text files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                text_content = await file.read()
            
            result['text'] = text_content
            result['metadata']['processing_method'] = 'plain_text'
            result['metadata']['lines'] = len(text_content.split('\n'))
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                async with aiofiles.open(file_path, 'r', encoding='latin-1') as file:
                    text_content = await file.read()
                
                result['text'] = text_content
                result['metadata']['processing_method'] = 'plain_text_latin1'
                result['metadata']['lines'] = len(text_content.split('\n'))
                
            except Exception as e:
                raise Exception(f"Text file processing failed: {str(e)}")
        
        return result
    
    async def process_uploaded_file(
        self, 
        file_content: bytes, 
        filename: str,
        use_vision: bool = False
    ) -> Dict[str, Any]:
        """
        Process uploaded file content
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            use_vision: Whether to use vision processing for PDFs
            
        Returns:
            Processing result dictionary
        """
        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as temp_file:
            temp_file.write(file_content)
            temp_path = temp_file.name
        
        try:
            result = await self.process_document(temp_path, use_vision=use_vision)
            result['metadata']['original_filename'] = filename
            return result
        finally:
            # Clean up temporary file
            try:
                os.unlink(temp_path)
            except Exception as e:
                logger.warning(f"Failed to clean up temp file {temp_path}: {str(e)}")
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return list(self.supported_formats)
    
    def get_capabilities(self) -> Dict[str, bool]:
        """Get processing capabilities"""
        return self.capabilities.copy()
    
    async def batch_process_documents(
        self, 
        file_paths: List[str],
        use_vision: bool = False
    ) -> List[Dict[str, Any]]:
        """
        Process multiple documents in batch
        
        Args:
            file_paths: List of file paths to process
            use_vision: Whether to use vision processing
            
        Returns:
            List of processing results
        """
        results = []
        
        for file_path in file_paths:
            try:
                result = await self.process_document(file_path, use_vision=use_vision)
                results.append(result)
            except Exception as e:
                logger.error(f"Batch processing failed for {file_path}: {str(e)}")
                results.append({
                    'text': '',
                    'metadata': {'file_path': file_path},
                    'success': False,
                    'error': str(e)
                })
        
        return results 